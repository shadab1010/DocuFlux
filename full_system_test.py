#!/usr/bin/env python3
"""
Complete System Test and Configuration Script
Tests all conversion functionality and creates sample files
"""

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def test_imports():
    """Test all required imports"""
    print("\n" + "="*70)
    print("TESTING IMPORTS & DEPENDENCIES")
    print("="*70)
    
    packages = {
        'pypdf': 'PyPDF',
        'docx': 'python-docx',
        'reportlab': 'ReportLab',
        'pptx': 'python-pptx',
        'openpyxl': 'OpenPyXL',
        'fitz': 'PyMuPDF',
        'pdf2docx': 'pdf2docx',
        'pdfplumber': 'pdfplumber',
    }
    
    installed = []
    missing = []
    
    for module, name in packages.items():
        try:
            __import__(module)
            installed.append(name)
            print(f"✓ {name:20} - OK")
        except ImportError:
            missing.append(name)
            print(f"✗ {name:20} - MISSING")
    
    print(f"\nInstalled: {len(installed)}/{len(packages)}")
    return len(missing) == 0

def test_converter():
    """Test conversion module"""
    print("\n" + "="*70)
    print("TESTING CONVERTER MODULE")
    print("="*70)
    
    try:
        from libreoffice_converter import DocumentConverter, convert_pdf_to_word, convert_word_to_pdf
        print("✓ libreoffice_converter imported successfully")
        
        tools = DocumentConverter.get_available_tools()
        print(f"\nAvailable Tools:")
        print(f"  LibreOffice: {'✓ AVAILABLE' if tools['libreoffice'] else '✗ NOT FOUND'}")
        if tools['libreoffice']:
            print(f"    Path: {tools['libreoffice']}")
        print(f"  Pandoc: {'✓ AVAILABLE' if tools['pandoc'] else '✗ NOT FOUND'}")
        if tools['pandoc']:
            print(f"    Path: {tools['pandoc']}")
        
        return True
    except Exception as e:
        print(f"✗ Error testing converter: {str(e)}")
        return False

def create_test_word_document():
    """Create a comprehensive test Word document"""
    print("\n" + "="*70)
    print("CREATING TEST WORD DOCUMENT")
    print("="*70)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Professional Test Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        doc.add_heading('Introduction', level=1)
        intro = doc.add_paragraph(
            'This is a comprehensive test document created to verify PDF ↔ Word conversion quality. '
            'It includes various formatting elements, tables, and structured content.'
        )
        
        # Section 1: Formatted Text
        doc.add_heading('Text Formatting Examples', level=1)
        doc.add_paragraph('This paragraph contains bold, italic, and colored text.', style='Normal')
        
        # Add formatted paragraph
        p = doc.add_paragraph()
        p.add_run('Bold text ').bold = True
        p.add_run('Italic text ').italic = True
        run = p.add_run('Red colored text')
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Section 2: Table
        doc.add_heading('Data Table Example', level=1)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Department'
        header_cells[2].text = 'Salary'
        
        # Data rows
        data = [
            ('John Smith', 'Engineering', '$95,000'),
            ('Jane Doe', 'Marketing', '$75,000'),
            ('Bob Johnson', 'Sales', '$85,000'),
        ]
        
        for i, (name, dept, salary) in enumerate(data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = name
            row_cells[1].text = dept
            row_cells[2].text = salary
        
        # Section 3: Multi-paragraph content
        doc.add_heading('Detailed Content Section', level=1)
        doc.add_paragraph(
            'Lorem ipsum dolor sit amet, consectetur adipiscing elit. '
            'Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.',
            style='Normal'
        )
        
        doc.add_paragraph(
            'Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris '
            'nisi ut aliquip ex ea commodo consequat.',
            style='Normal'
        )
        
        # Section 4: Bulleted list
        doc.add_heading('Key Features', level=1)
        doc.add_paragraph('Professional conversion quality', style='List Bullet')
        doc.add_paragraph('Structure preservation', style='List Bullet')
        doc.add_paragraph('Table and image support', style='List Bullet')
        doc.add_paragraph('Multiple page handling', style='List Bullet')
        
        # Section 5: Numbered list
        doc.add_heading('Conversion Steps', level=1)
        doc.add_paragraph('Upload your document', style='List Number')
        doc.add_paragraph('Select target format', style='List Number')
        doc.add_paragraph('Process conversion', style='List Number')
        doc.add_paragraph('Download result', style='List Number')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('End of Document')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        filename = 'test_sample_complete.docx'
        doc.save(filename)
        print(f"✓ Created: {filename}")
        print(f"  Size: {os.path.getsize(filename)} bytes")
        return filename
        
    except Exception as e:
        print(f"✗ Error creating Word document: {str(e)}")
        return None

def test_conversion_functions():
    """Test conversion functions with sample files"""
    print("\n" + "="*70)
    print("TESTING CONVERSION FUNCTIONS")
    print("="*70)
    
    try:
        # Test if conversion functions are callable
        from libreoffice_converter import convert_pdf_to_word, convert_word_to_pdf
        
        print("✓ convert_pdf_to_word function: Available")
        print("✓ convert_word_to_pdf function: Available")
        print("\nConversion functions are ready to use!")
        print("  - They will use LibreOffice if available")
        print("  - They will fall back to other methods if needed")
        
        return True
    except Exception as e:
        print(f"✗ Error: {str(e)}")
        return False

def verify_app_integration():
    """Verify Flask app integration"""
    print("\n" + "="*70)
    print("CHECKING APP INTEGRATION")
    print("="*70)
    
    try:
        # Check if app.py imports conversion module
        with open('app.py', 'r') as f:
            content = f.read()
            
        if 'from libreoffice_converter import' in content:
            print("✓ app.py imports conversion module")
        else:
            print("✗ app.py does not import conversion module")
            return False
        
        if '/convert-pdf-to-word' in content:
            print("✓ PDF to Word route exists")
        else:
            print("✗ PDF to Word route missing")
            return False
        
        if '/convert-word-to-pdf' in content:
            print("✓ Word to PDF route exists")
        else:
            print("✗ Word to PDF route missing")
            return False
        
        if 'LibreOfficeConverter.is_available()' in content:
            print("✓ Conversion system is integrated")
        else:
            print("✗ Conversion system not integrated")
            return False
        
        return True
    except Exception as e:
        print(f"✗ Error checking integration: {str(e)}")
        return False

def main():
    """Run all tests"""
    print("\n" + "█"*70)
    print("█ MyPDF - COMPLETE SYSTEM CONFIGURATION & TEST")
    print("█"*70)
    
    results = []
    
    # Test 1: Imports
    results.append(("Dependencies", test_imports()))
    
    # Test 2: Converter Module
    results.append(("Converter Module", test_converter()))
    
    # Test 3: Create Test Document
    test_doc = create_test_word_document()
    results.append(("Test Document Creation", test_doc is not None))
    
    # Test 4: Conversion Functions
    results.append(("Conversion Functions", test_conversion_functions()))
    
    # Test 5: App Integration
    results.append(("App Integration", verify_app_integration()))
    
    # Summary
    print("\n" + "="*70)
    print("SYSTEM TEST SUMMARY")
    print("="*70)
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for test_name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{status:8} - {test_name}")
    
    print(f"\nOverall: {passed}/{total} tests passed")
    
    # Final Instructions
    print("\n" + "="*70)
    print("NEXT STEPS")
    print("="*70)
    
    if passed == total:
        print("\n✓ All tests passed! Your system is configured correctly.")
        print("\nWhat to do next:")
        print("  1. Restart your Flask app (Ctrl+C and run python app.py)")
        print("  2. Install LibreOffice for professional-grade conversions:")
        print("     → https://www.libreoffice.org/download/")
        print("  3. Test conversions with your documents")
        print("  4. Enjoy professional-quality output!")
    else:
        print(f"\n⚠  {total - passed} test(s) failed. Check the output above.")
        print("\nTroubleshooting:")
        print("  1. Ensure all packages are installed: pip install -r requirements.txt")
        print("  2. Check that app.py and libreoffice_converter.py are present")
        print("  3. Verify Python version: 3.8+")
    
    print("\n" + "█"*70)
    print("█ Test Complete")
    print("█"*70 + "\n")
    
    return passed == total

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)

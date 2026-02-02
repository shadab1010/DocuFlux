#!/usr/bin/env python
"""
Test script to verify conversion functionality
Run this to test PDF<->Word conversions
"""

import os
import sys
from pathlib import Path

# Add project to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from libreoffice_converter import DocumentConverter, convert_pdf_to_word, convert_word_to_pdf

def test_available_tools():
    """Test what conversion tools are available"""
    print("=" * 60)
    print("CHECKING AVAILABLE CONVERSION TOOLS")
    print("=" * 60)
    
    tools = DocumentConverter.get_available_tools()
    
    print(f"\nLibreOffice: {'✓ AVAILABLE' if tools['libreoffice'] else '✗ NOT FOUND'}")
    if tools['libreoffice']:
        print(f"  Path: {tools['libreoffice']}")
    else:
        print("  Action: Install from https://www.libreoffice.org/download/")
    
    print(f"\nPandoc: {'✓ AVAILABLE' if tools['pandoc'] else '✗ NOT FOUND'}")
    if tools['pandoc']:
        print(f"  Path: {tools['pandoc']}")
    
    return tools

def create_test_files():
    """Create sample test files"""
    print("\n" + "=" * 60)
    print("CREATING TEST FILES")
    print("=" * 60)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create test Word document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_heading('Heading 1', level=1)
        doc.add_paragraph('This is a test paragraph with some text.')
        doc.add_paragraph('Another paragraph with ', style='Normal')
        
        # Add table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        table.cell(2, 0).text = 'Data 3'
        table.cell(2, 1).text = 'Data 4'
        
        test_word = 'test_document.docx'
        doc.save(test_word)
        print(f"✓ Created: {test_word}")
        
        return test_word
        
    except Exception as e:
        print(f"✗ Error creating test document: {str(e)}")
        return None

def main():
    """Run all tests"""
    print("\n" + "=" * 60)
    print("MyPDF CONVERSION SYSTEM TEST")
    print("=" * 60)
    
    # 1. Check available tools
    tools = test_available_tools()
    
    # 2. Create test files
    test_file = create_test_files()
    
    if test_file:
        print(f"\nTest document created: {test_file}")
        print("\nYou can now test conversions:")
        print(f"  1. Test Word→PDF: Convert {test_file} to PDF")
        print(f"  2. Upload to your web app")
        print(f"  3. Check quality in output files")
    
    print("\n" + "=" * 60)
    print("RECOMMENDATIONS")
    print("=" * 60)
    
    if not tools['libreoffice']:
        print("\n⚠️  LibreOffice NOT INSTALLED")
        print("\nTo get best quality conversions:")
        print("1. Visit: https://www.libreoffice.org/download/")
        print("2. Download Windows installer")
        print("3. Install with default settings")
        print("4. Restart your Flask app")
        print("\nWithout LibreOffice, conversions will use fallback methods")
    else:
        print("\n✓ LibreOffice is installed - you'll get professional-grade conversions!")
    
    print("\n" + "=" * 60)
    print("TEST COMPLETE")
    print("=" * 60)

if __name__ == '__main__':
    main()

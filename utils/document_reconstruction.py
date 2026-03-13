import fitz
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import io
from collections import Counter
import re

class DocumentReconstructionEngine:
    def __init__(self, pdf_path):
        self.pdf_path = pdf_path
        self.doc = fitz.open(pdf_path)
        self.docx = Document()
        self.base_font_size = 11
        self.extracted_images = []

    def _calculate_base_font(self):
        sizes = []
        for page in self.doc:
            blocks = page.get_text("dict").get("blocks", [])
            for b in blocks:
                if b.get("type") == 0:  # Text block
                    for line in b.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                sizes.append(round(span.get("size", 11)))
        if sizes:
            most_common = Counter(sizes).most_common(1)
            if most_common:
                self.base_font_size = most_common[0][0]

    def _is_code_block(self, block):
        # Heuristic for code blocks: monospace font or starts with common code indentation
        # Check spans for typical monospace font names
        mono_fonts = ['Courier', 'Consolas', 'Mono', 'Andale', 'Terminal']
        lines = block.get("lines", [])
        code_score = 0
        total_spans = 0
        
        for line in lines:
            for span in line.get("spans", []):
                total_spans += 1
                font_name = span.get("font", "")
                if any(m.lower() in font_name.lower() for m in mono_fonts):
                    code_score += 1
        
        if total_spans > 0 and (code_score / total_spans) >= 0.5:
            return True
            
        return False

    def _is_bullet_list(self, line_text):
        bullet_chars = ['•', '', 'v', 'o', '-', '*', '➢', '']
        strip_text = line_text.strip()
        if not strip_text: return False
        if strip_text[0] in bullet_chars:
            return True
        return False

    def _get_heading_level(self, size, is_bold):
        diff = size - self.base_font_size
        if diff >= 8: return 1
        if diff >= 4: return 2
        if diff >= 2: return 3
        if is_bold and diff >= 0: return 4
        return 0

    def _is_numbered_list(self, line_text):
        """Detect 1. 2. a. b. i. ii. style list items"""
        strip_text = line_text.strip()
        if re.match(r'^(\d+\.|[a-zA-Z]\.|[ivxIVX]+\.)[\s]', strip_text):
            return True
        return False

    def _add_code_shading(self, paragraph):
        """Add light grey background shading to a paragraph for code blocks"""
        try:
            pPr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F0F0F0')  # light grey
            pPr.append(shd)
        except Exception:
            pass

    def _set_paragraph_spacing(self, paragraph, space_before=0, space_after=6):
        """Set space before/after a paragraph in points"""
        try:
            pf = paragraph.paragraph_format
            if space_before > 0:
                pf.space_before = Pt(space_before)
            if space_after > 0:
                pf.space_after = Pt(space_after)
        except Exception:
            pass


    def _extract_page_images(self, page):
        img_info = []
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                base_image = self.doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_dict = {
                    "bytes": image_bytes,
                    "ext": base_image["ext"],
                    "y0": 0  # We need to find the approximate y0 position
                }
                
                # Try to find exactly where this image is placed
                rects = page.get_image_rects(xref)
                if rects:
                    img_dict["y0"] = rects[0].y0
                img_info.append(img_dict)
            except:
                pass
        return img_info

    def _hex_to_rgb(self, srg_color):
        if not srg_color: return RGBColor(0, 0, 0)
        # srg_color is usually an integer representing RGB
        b = srg_color & 255
        g = (srg_color >> 8) & 255
        r = (srg_color >> 16) & 255
        return RGBColor(r, g, b)

    def _normalize_font_name(self, fontname):
        if not fontname: return 'Calibri'
        fontname_lower = fontname.lower()
        if 'times' in fontname_lower: return 'Times New Roman'
        if 'courier' in fontname_lower: return 'Courier New'
        if 'helvetica' in fontname_lower: return 'Arial'
        if 'arial' in fontname_lower: return 'Arial'
        if 'consolas' in fontname_lower: return 'Consolas'
        
        if '+' in fontname: fontname = fontname.split('+')[1]
        for suffix in ['-Bold', '-Italic', '-Regular', ',Bold', ',Italic']:
            if fontname.endswith(suffix):
                fontname = fontname[:-len(suffix)]
        return fontname or 'Calibri'

    def _group_into_columns(self, blocks):
        """Detect multi-column aligned labels / content by grouping y overlaps and x gaps"""
        # Sort vertically
        blocks.sort(key=lambda b: b.get("bbox", [0, 0, 0, 0])[1])
        
        rows = []
        current_row = []
        row_y = -1
        y_tolerance = 15 # Points of vertical overlap tolerance
        
        for block in blocks:
            bbox = block.get("bbox", [0, 0, 0, 0])
            y0, y1 = bbox[1], bbox[3]
            
            if row_y == -1 or abs(y0 - row_y) <= y_tolerance:
                current_row.append(block)
                if row_y == -1: row_y = y0
            else:
                rows.append(current_row)
                current_row = [block]
                row_y = y0
                
        if current_row:
            rows.append(current_row)
            
        final_structures = []
        # For each row, check if there are distinct X columns
        for row in rows:
            if len(row) > 1:
                # Sort row blocks horizontally
                row.sort(key=lambda b: b.get("bbox", [0])[0])
                # We have multiple blocks on the same vertical line, this is a table/column
                final_structures.append({
                    "type": "columns",
                    "blocks": row,
                    "y0": row[0].get("bbox")[1]
                })
            elif len(row) == 1:
                final_structures.append({
                    "type": "single",
                    "block": row[0],
                    "y0": row[0].get("bbox")[1]
                })
                
        return final_structures

    def _render_spans_to_paragraph(self, lines, paragraph, is_code=False):
        full_text = ""
        for line in lines:
            line_text = ""
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text: continue
                line_text += text
                
                run = paragraph.add_run(text)
                
                if is_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    run.font.name = self._normalize_font_name(span.get("font", ""))
                    run.font.size = Pt(round(span.get("size", 11)))
                    # Check styles
                    flags = span.get("flags", 0)
                    if flags & 2**4: run.font.bold = True
                    if flags & 2**1: run.font.italic = True
                    run.font.color.rgb = self._hex_to_rgb(span.get("color", 0))
            
            # If not last line in the block, add a return or space
            if line != lines[-1]:
                if is_code:
                    paragraph.add_run('\n')
                else:
                    paragraph.add_run(' ')
            full_text += line_text + "\n"
        return full_text

    def _process_page(self, page):
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        
        # separate text blocks and non-text blocks
        text_blocks = [b for b in blocks if b.get("type") == 0]
        
        # Images handling
        images = self._extract_page_images(page)
        
        # Structural layout grouping
        structures = self._group_into_columns(text_blocks)
        
        # Mix images into structures based on y0
        for img in images:
            structures.append({
                "type": "image",
                "bytes": img["bytes"],
                "ext": img["ext"],
                "y0": img["y0"]
            })
            
        structures.sort(key=lambda s: s.get("y0", 0))
        
        for struct in structures:
            stype = struct.get("type")
            
            if stype == "image":
                try:
                    img_stream = io.BytesIO(struct["bytes"])
                    self.docx.add_picture(img_stream, width=Inches(6.0))
                except Exception as e:
                    print(f"Error adding image: {e}")
                    
            elif stype == "columns":
                # Create a table for columns
                cols = struct.get("blocks", [])
                table = self.docx.add_table(rows=1, cols=len(cols))
                # remove borders
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                
                for idx, col_block in enumerate(cols):
                    cell = table.cell(0, idx)
                    # clear default paragraph
                    p = cell.paragraphs[0]
                    self._render_spans_to_paragraph(col_block.get("lines", []), p)
                    
            elif stype == "single":
                block = struct.get("block")
                lines = block.get("lines", [])
                if not lines: continue
                         # Check for code blocks
                if self._is_code_block(block):
                    p = self.docx.add_paragraph()
                    self._set_paragraph_spacing(p, space_before=4, space_after=4)
                    self._render_spans_to_paragraph(lines, p, is_code=True)
                    self._add_code_shading(p)
                    continue

                # Get first span properties to determine Heading Level or List
                first_span = None
                for line in lines:
                    if line.get("spans"):
                        first_span = line["spans"][0]
                        break
                        
                # Check bullet list
                first_text = ""
                if first_span: first_text = first_span.get("text", "").strip()

                if first_text and self._is_bullet_list(first_text):
                    p = self.docx.add_paragraph(style='List Bullet')
                    # strip the bullet character and leading space from first span
                    lines_copy = []
                    for il, line in enumerate(lines):
                        new_line = {"spans": []}
                        for isp, span in enumerate(line.get("spans", [])):
                            text = span.get("text", "")
                            if il == 0 and isp == 0:
                                text = re.sub(r'^[•·‣‒–vo\-*➢▪▸○]\s*', '', text)
                            new_span = dict(span)
                            new_span["text"] = text
                            new_line["spans"].append(new_span)
                        lines_copy.append(new_line)
                    self._render_spans_to_paragraph(lines_copy, p)
                    continue

                if first_text and self._is_numbered_list(first_text):
                    p = self.docx.add_paragraph(style='List Number')
                    # Strip the leading number/letter label from first span
                    lines_copy = []
                    for il, line in enumerate(lines):
                        new_line = {"spans": []}
                        for isp, span in enumerate(line.get("spans", [])):
                            text = span.get("text", "")
                            if il == 0 and isp == 0:
                                text = re.sub(r'^(\d+\.|[a-zA-Z]\.|[ivxIVX]+\.)\s*', '', text)
                            new_span = dict(span)
                            new_span["text"] = text
                            new_line["spans"].append(new_span)
                        lines_copy.append(new_line)
                    self._render_spans_to_paragraph(lines_copy, p)
                    continue
                
                # Check heading
                p = self.docx.add_paragraph()
                if first_span:
                    size = first_span.get("size", 11)
                    flags = first_span.get("flags", 0)
                    is_bold = bool(flags & 2**4)
                    heading_level = self._get_heading_level(size, is_bold)
                    
                    if heading_level > 0:
                        try:
                            p.style = f'Heading {min(heading_level, 9)}'
                            self._set_paragraph_spacing(p, space_before=10, space_after=4)
                        except:
                            pass  # default to normal if style missing
                    else:
                        self._set_paragraph_spacing(p, space_before=0, space_after=4)

                self._render_spans_to_paragraph(lines, p)

    def convert(self, output_path):
        try:
            self._calculate_base_font()
            for page_num in range(len(self.doc)):
                page = self.doc[page_num]
                self._process_page(page)
                if page_num < len(self.doc) - 1:
                    self.docx.add_page_break()

            self.docx.save(output_path)
            self.doc.close()
            return True, output_path
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, str(e)

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 2:
        engine = DocumentReconstructionEngine(sys.argv[1])
        engine.convert(sys.argv[2])
        print("Success")

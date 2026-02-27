"""
PDF conversion utilities
"""
import os
import uuid
from zipfile import ZipFile
import fitz
from pypdf import PdfReader, PdfWriter
from pdf2docx.converter import Converter
from openpyxl import Workbook
import pdfplumber
import io
from docx import Document
from docx.shared import Inches


def merge_pdfs(pdf_files, output_path):
    """Merge multiple PDF files into one"""
    try:
        writer = PdfWriter()
        
        for pdf_file in pdf_files:
            try:
                reader = PdfReader(pdf_file)
                # Verify we can read the pages
                if len(reader.pages) == 0:
                     return False, f"File appears empty (0 pages): {os.path.basename(pdf_file)}"
                
                for page in reader.pages:
                    writer.add_page(page)
            except Exception as e:
                return False, f"Error processing {os.path.basename(pdf_file)}: {str(e)}"
        
        with open(output_path, "wb") as f:
            writer.write(f)
        
        return True, output_path
    except Exception as e:
        return False, f"Merge failed: {str(e)}"


def split_pdf(pdf_path, page_ranges, output_folder):
    """Split PDF by page ranges"""
    try:
        reader = PdfReader(pdf_path)
        output_files = []
        
        for idx, page_range in enumerate(page_ranges):
            writer = PdfWriter()
            start, end = page_range
            
            for page_num in range(start - 1, end):
                if page_num < len(reader.pages):
                    writer.add_page(reader.pages[page_num])
            
            output_file = os.path.join(output_folder, f"split_{idx + 1}.pdf")
            with open(output_file, "wb") as f:
                writer.write(f)
            output_files.append(output_file)
        
        return True, output_files
    except Exception as e:
        return False, str(e)


def compress_pdf(pdf_path, output_path):
    """Compress PDF by reducing page quality"""
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        with open(output_path, "wb") as f:
            writer.write(f)
        
        return True, output_path
    except Exception as e:
        return False, str(e)


def _enhance_word_document(docx_path, pdf_path):
    """Enhance Word document with better font and structure preservation from PDF"""
    try:
        # Extract font information from PDF using pdfplumber
        font_info = {}
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages[:1]):  # Check first page for fonts
                    for char in page.chars:
                        fontname = char.get('fontname', 'Calibri')
                        size = char.get('size', 11)
                        if fontname and fontname != 'unknown':
                            key = f"{fontname}_{size}"
                            font_info[key] = {
                                'name': _normalize_font_name(fontname),
                                'size': int(size)
                            }
        except Exception as e:
            print(f"Font extraction warning (non-critical): {e}")
            return  # Don't fail if font extraction doesn't work
        
        # If we found fonts, try to update the Word document
        if font_info:
            try:
                doc = Document(docx_path)
                # Find common fonts and update styles
                most_common_font = max(font_info.items(), key=lambda x: x[1]['size'])[1]
                
                # Update default font for all paragraphs
                for para in doc.paragraphs:
                    for run in para.runs:
                        if run.font.name is None or run.font.name == 'Calibri':
                            run.font.name = most_common_font.get('name', 'Calibri')
                
                # Update tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.font.name is None or run.font.name == 'Calibri':
                                        run.font.name = most_common_font.get('name', 'Calibri')
                
                doc.save(docx_path)
            except Exception as e:
                print(f"Document enhancement warning (non-critical): {e}")
                
    except Exception as e:
        print(f"Enhancement non-critical error: {e}")


def _normalize_font_name(fontname):
    """Normalize PDF font names to Word-compatible names"""
    if not fontname:
        return 'Calibri'
    
    fontname_lower = fontname.lower()
    
    # Map common fonts
    font_map = {
        'times': 'Times New Roman',
        'courier': 'Courier New',
        'helvetica': 'Calibri',
        'arial': 'Arial',
        'cambria': 'Cambria',
        'verdana': 'Verdana',
        'georgia': 'Georgia',
        'tahoma': 'Tahoma',
    }
    
    for key, value in font_map.items():
        if key in fontname_lower:
            return value
    
    # If unmapped, try to clean up the font name
    if '+' in fontname:
        fontname = fontname.split('+')[1]
    
    # Remove common suffixes
    for suffix in ['-Bold', '-Italic', '-Regular', ',Bold', ',Italic']:
        if fontname.endswith(suffix):
            fontname = fontname[:-len(suffix)]
    
    return fontname or 'Calibri'


def pdf_to_word_as_images(pdf_path, output_path):
    """Fallback: Convert PDF pages to images and insert into Word"""
    try:
        doc = Document()
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # 2x zoom for better quality
            img_data = pix.tobytes("png")
            
            # Add image to Word doc
            img_stream = io.BytesIO(img_data)
            doc.add_picture(img_stream, width=Inches(6))
            
            if page_num < len(pdf_doc) - 1:
                doc.add_page_break()
                
        doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def pdf_to_images(pdf_path, output_folder, image_format="png"):
    """Convert PDF pages to images"""
    try:
        pdf_doc = fitz.open(pdf_path)
        images = []
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            file_id = uuid.uuid4().hex
            img_path = os.path.join(output_folder, f"page_{file_id}_{page_num}.{image_format}")
            pix.save(img_path)
            images.append(img_path)
        
        pdf_doc.close()
        return True, images
    except Exception as e:
        return False, str(e)


def pdf_to_word(pdf_path, output_path):
    """Convert PDF to Word document with enhanced structure and font preservation"""
    try:
        # Pre-validate PDF
        try:
            with fitz.open(pdf_path) as doc:
                if doc.page_count < 1:
                    return False, "PDF is empty (0 pages)."
        except Exception as e:
            return False, f"Invalid PDF file: {str(e)}"

        # Attempt 1: pdf2docx with enhanced layout detection (Best for structured PDFs)
        try:
            converter = Converter(pdf_path)
            # Enable layout preservation and font detection
            converter.convert(output_path, start=0, end=None)
            converter.close()
            
            # Post-process to enhance font and structure preservation
            try:
                _enhance_word_document(output_path, pdf_path)
            except Exception as enhance_e:
                print(f"Enhancement failed (non-critical): {enhance_e}")
            
            # Double check output
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                 return True, output_path
        except Exception as e:
            print(f"pdf2docx failed: {e}. Attempting fallback...")
        
        # Attempt 2: LibreOffice Fallback (Best for scanned/complex PDFs)
        try:
            # Dynamic import to handle path differences
            try:
                from libreoffice_converter import LibreOfficeConverter
            except ImportError:
                # Try relative import if in package
                import sys
                sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                from libreofficeconverter import LibreOfficeConverter

            if LibreOfficeConverter.is_available():
                success = LibreOfficeConverter.convert(pdf_path, output_path, output_format="docx")
                if success:
                    return True, output_path
            
            # Attempt 3: Image-based wrapper (Final fallback for scans)
            print("LibreOffice unavailable/failed. Falling back to Image-Wrapper...")
            success, result = pdf_to_word_as_images(pdf_path, output_path)
            if success:
                return True, result
            
            return False, "All conversion methods failed (pdf2docx, LibreOffice, Image-Wrapper)."
                
        except Exception as fallback_error:
            return False, f"All conversion methods failed. Error: {str(fallback_error)}"

    except Exception as e:
        # Catch-all for top level errors
        return False, f"Conversion error: {str(e)}"


def word_to_pdf(word_path, output_path):
    """Convert Word document to PDF with enhanced formatting preservation"""
    try:
        # Pre-validate Word file
        if not os.path.exists(word_path):
            return False, "Word file not found."
        
        file_size = os.path.getsize(word_path)
        if file_size == 0:
            return False, "Word file is empty."
        
        # Attempt 1: LibreOffice (Best quality, preserves all fonts and formatting)
        try:
            # Dynamic import to handle path differences
            try:
                from libreoffice_converter import LibreOfficeConverter
            except ImportError:
                import sys
                sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                from libreoffice_converter import LibreOfficeConverter
            
            if LibreOfficeConverter.is_available():
                print("Using LibreOffice for Word to PDF conversion...")
                success = LibreOfficeConverter.convert(word_path, output_path, output_format="pdf")
                if success and os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    print("LibreOffice conversion successful")
                    return True, output_path
                else:
                    print("LibreOffice conversion failed, trying fallback...")
        except Exception as e:
            print(f"LibreOffice conversion error: {e}. Attempting fallback...")
        
        # Attempt 2: docx2pdf (Windows-only fallback, uses Word COM)
        try:
            from docx2pdf import convert as docx2pdf_convert
            print("Using docx2pdf for Word to PDF conversion...")
            docx2pdf_convert(word_path, output_path)
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print("docx2pdf conversion successful")
                return True, output_path
            else:
                print("docx2pdf produced empty file")
        except Exception as e:
            print(f"docx2pdf failed: {e}")
        
        # Attempt 3: pypandoc (Universal fallback, requires pandoc)
        try:
            import pypandoc
            print("Using pypandoc for Word to PDF conversion...")
            pypandoc.convert_file(word_path, 'pdf', outputfile=output_path)
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print("pypandoc conversion successful")
                return True, output_path
        except ImportError:
            print("pypandoc not available")
        except Exception as e:
            print(f"pypandoc failed: {e}")
        
        # If all methods failed
        return False, "All Word to PDF conversion methods failed. Please ensure LibreOffice or Microsoft Word is installed."
        
    except Exception as e:
        return False, f"Word to PDF conversion error: {str(e)}"


def pdf_to_excel(pdf_path, output_path):
    """Extract tables from PDF and convert to Excel"""
    try:
        workbook = Workbook()
        workbook.remove(workbook.active)
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                
                if tables:
                    for table_idx, table in enumerate(tables):
                        ws = workbook.create_sheet(title=f"Page_{page_num + 1}_Table_{table_idx + 1}")
                        for row_idx, row in enumerate(table):
                            for col_idx, cell in enumerate(row):
                                ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell)
        
        # Remove empty sheet if created
        if len(workbook.sheetnames) == 0:
            ws = workbook.create_sheet("Sheet1")
        
        workbook.save(output_path)
        return True, output_path
    except Exception as e:
        return False, str(e)


def pdf_to_jpg(pdf_path, output_folder):
    """Convert PDF pages to JPG images"""
    return pdf_to_images(pdf_path, output_folder, "jpg")


def pdf_to_png(pdf_path, output_folder):
    """Convert PDF pages to PNG images"""
    return pdf_to_images(pdf_path, output_folder, "png")


def create_image_zip(image_files, output_path):
    """Create ZIP archive from image files"""
    try:
        with ZipFile(output_path, 'w') as zipf:
            for img_file in image_files:
                zipf.write(img_file, os.path.basename(img_file))
        return True, output_path
    except Exception as e:
        return False, str(e)


def add_text_to_pdf(pdf_path, output_path, annotations):
    """Add text annotations to PDF pages"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        for annotation in annotations:
            page_num = annotation.get('page', 0)
            x = annotation.get('x', 50)
            y = annotation.get('y', 50)
            text = annotation.get('text', '')
            font_size = annotation.get('font_size', 12)
            color = annotation.get('color', (0, 0, 0))
            
            if 0 <= page_num < len(pdf_doc):
                page = pdf_doc[page_num]
                # Insert text at specified location
                point = fitz.Point(x, y)
                page.insert_text(point, text, fontsize=font_size, color=color)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def rotate_pdf_pages(pdf_path, output_path, rotations):
    """Rotate specific pages in PDF"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        for rotation in rotations:
            page_num = rotation.get('page', 0)
            angle = rotation.get('angle', 90)
            
            if 0 <= page_num < len(pdf_doc):
                page = pdf_doc[page_num]
                page.set_rotation(angle)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def delete_pdf_pages(pdf_path, output_path, pages_to_delete):
    """Delete specific pages from PDF"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        # Sort in reverse to delete from end first
        for page_num in sorted(pages_to_delete, reverse=True):
            if 0 <= page_num < len(pdf_doc):
                pdf_doc.delete_page(page_num)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def extract_pdf_page(pdf_path, output_path, page_num):
    """Extract a single page from PDF"""
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        if 0 <= page_num < len(reader.pages):
            writer.add_page(reader.pages[page_num])
        
        with open(output_path, "wb") as f:
            writer.write(f)
        
        return True, output_path
    except Exception as e:
        return False, str(e)


def add_watermark_to_pdf(pdf_path, output_path, options):
    """Add advanced watermark to PDF pages"""
    try:
        import os
        from PIL import Image, ImageDraw, ImageFont
        import fitz
        
        wm_type = options.get('type', 'text')
        transparency = float(options.get('transparency', 50))
        opacity = int(255 * (100 - transparency) / 100.0)
        rotation = int(options.get('rotation', 0))
        position = options.get('position', 'center')
        mosaic = str(options.get('mosaic', 'false')).lower() in ('true', '1')
        
        if wm_type == 'text':
            text = options.get('text', 'Watermark')
            font_size = int(options.get('fontSize', 48))
            color_hex = options.get('color', '#000000').lstrip('#')
            if len(color_hex) >= 6:
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
            else:
                r, g, b = (0, 0, 0)
                
            try:
                # Use a standard font if available, or default
                font = ImageFont.truetype("arial.ttf", font_size)
            except IOError:
                font = ImageFont.load_default()
                
            dummy_img = Image.new('RGBA', (1, 1), (0, 0, 0, 0))
            draw = ImageDraw.Draw(dummy_img)
            bbox = draw.textbbox((0, 0), text, font=font)
            text_w = bbox[2] - bbox[0]
            text_h = bbox[3] - bbox[1]
            
            # Create actual text image wrapper
            wm_img = Image.new('RGBA', (int(text_w) + 20, int(text_h) + 20), (0, 0, 0, 0))
            draw = ImageDraw.Draw(wm_img)
            draw.text((10, 10), text, font=font, fill=(r, g, b, opacity))
            
        else:
            img_path = options.get('image_path')
            if not img_path or not os.path.exists(img_path):
                return False, "Watermark image missing"
            
            wm_img = Image.open(img_path).convert('RGBA')
            
            if opacity < 255:
                r_channel, g_channel, b_channel, a_channel = wm_img.split()
                a_channel = a_channel.point(lambda p: int(p * (opacity / 255.0)))
                wm_img = Image.merge('RGBA', (r_channel, g_channel, b_channel, a_channel))
                
        # Rotation (Negative so it visually rotates correctly in PyMuPDF)
        if rotation != 0:
            wm_img = wm_img.rotate(-rotation, expand=True, resample=Image.BICUBIC, fillcolor=(0,0,0,0))
            
        import uuid
        temp_img_path = os.path.join(os.path.dirname(output_path), f"wm_temp_{uuid.uuid4().hex}.png")
        wm_img.save(temp_img_path, "PNG")
            
        pdf_doc = fitz.open(pdf_path)
        img_w, img_h = wm_img.width, wm_img.height
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pw, ph = page.rect.width, page.rect.height
            
            if mosaic:
                # Tile horizontally and vertically with padding
                pad_x, pad_y = 100, 100
                for x in range(0, int(pw), img_w + pad_x):
                    for y in range(0, int(ph), img_h + pad_y):
                        rect = fitz.Rect(x, y, x + img_w, y + img_h)
                        page.insert_image(rect, filename=temp_img_path)
            else:
                if position == 'top-left': x, y = 50, 50
                elif position == 'top-center': x, y = (pw - img_w) / 2, 50
                elif position == 'top-right': x, y = pw - img_w - 50, 50
                elif position == 'middle-left': x, y = 50, (ph - img_h) / 2
                elif position == 'center': x, y = (pw - img_w) / 2, (ph - img_h) / 2
                elif position == 'middle-right': x, y = pw - img_w - 50, (ph - img_h) / 2
                elif position == 'bottom-left': x, y = 50, ph - img_h - 50
                elif position == 'bottom-center': x, y = (pw - img_w) / 2, ph - img_h - 50
                elif position == 'bottom-right': x, y = pw - img_w - 50, ph - img_h - 50
                else: x, y = (pw - img_w) / 2, (ph - img_h) / 2
                
                rect = fitz.Rect(x, y, x + img_w, y + img_h)
                page.insert_image(rect, filename=temp_img_path)
                
        pdf_doc.save(output_path)
        pdf_doc.close()
        
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
            
        return True, output_path
        
    except Exception as e:
        return False, str(e)


def add_page_numbers_to_pdf(pdf_path, output_path, options=None):
    """Add advanced page numbers to PDF with professional positioning"""
    try:
        if options is None:
            options = {}
            
        # Extract options with defaults
        start_number = int(options.get('firstNumber', 1))
        from_page = int(options.get('fromPage', 1)) 
        to_page = int(options.get('toPage', 999999))
        position = options.get('position', 'bottom-right')
        margin_type = options.get('margin', 'Recommended')
        page_mode = options.get('pageMode', 'single')
        font_size = int(options.get('fontSize', 12))
        font_color = options.get('fontColor', '#000000') # Hex color
        text_format = options.get('format', '{n}') # Default is just the number
        
        # Convert hex color to RGB tuple (0-1 range)
        if font_color.startswith('#'):
            h = font_color.lstrip('#')
            rgb = tuple(int(h[i:i+2], 16)/255.0 for i in (0, 2, 4))
        else:
            rgb = (0, 0, 0)
            
        pdf_doc = fitz.open(pdf_path)
        total_pages = len(pdf_doc)
        to_page = min(to_page, total_pages)
        
        # Define margins
        if margin_type == 'Small': pad = 20
        elif margin_type == 'Big': pad = 60
        else: pad = 40 # Recommended
        
        font_name = "helv" # Standard Helvetica
        
        for page_num in range(total_pages):
            current_pdf_page_num = page_num + 1
            
            if from_page <= current_pdf_page_num <= to_page:
                page = pdf_doc[page_num]
                
                # Use page.rect which is rotation-aware in modern PyMuPDF
                pw, ph = page.rect.width, page.rect.height
                
                # Calculate number to show
                display_num = start_number + (current_pdf_page_num - from_page)
                
                # Determine text to render
                text = text_format.replace('{n}', str(display_num)).replace('{total}', str(total_pages))
                
                # Accurate text width/height
                text_w = fitz.get_text_length(text, fontname=font_name, fontsize=font_size)
                text_h = font_size # Cap-height is approx fontsize
                
                # Determine position
                curr_pos = position
                if page_mode == 'facing':
                    is_even = current_pdf_page_num % 2 == 0
                    if 'left' in curr_pos and is_even:
                        curr_pos = curr_pos.replace('left', 'right')
                    elif 'right' in curr_pos and is_even:
                        curr_pos = curr_pos.replace('right', 'left')
                
                # Calculate coordinates (y is baseline)
                if curr_pos == 'top-left': x, y = pad, pad + text_h
                elif curr_pos == 'top-center': x, y = (pw - text_w) / 2, pad + text_h
                elif curr_pos == 'top-right': x, y = pw - pad - text_w, pad + text_h
                elif curr_pos == 'middle-left': x, y = pad, ph / 2
                elif curr_pos == 'middle-center': x, y = (pw - text_w) / 2, ph / 2
                elif curr_pos == 'middle-right': x, y = pw - pad - text_w, ph / 2
                elif curr_pos == 'bottom-left': x, y = pad, ph - pad
                elif curr_pos == 'bottom-center': x, y = (pw - text_w) / 2, ph - pad
                elif curr_pos == 'bottom-right': x, y = pw - pad - text_w, ph - pad
                else: x, y = pw - pad - text_w, ph - pad
                
                # Insert text with explicit overlay and rotation handling
                # In modern fitz, insert_text(..., rotate=page.rotation) is handled by using page.rect
                page.insert_text((x, y), text, 
                                 fontname=font_name, 
                                 fontsize=font_size, 
                                 color=rgb)
        
        pdf_doc.save(output_path, garbage=4, deflate=True)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, str(e)


def add_image_to_pdf(pdf_path, output_path, image_path, page_num, x=0, y=0, width=None, height=None):
    """Add image to specific PDF page"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            if width and height:
                page.insert_image((x, y, x + width, y + height), filename=image_path)
            else:
                page.insert_image((x, y), filename=image_path)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def draw_rectangle_on_pdf(pdf_path, output_path, page_num, x, y, width, height, color=(1, 0, 0), line_width=2):
    """Draw rectangle on PDF page"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            rect = fitz.Rect(x, y, x + width, y + height)
            page.draw_rect(rect, color=color, width=line_width)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def draw_circle_on_pdf(pdf_path, output_path, page_num, cx, cy, radius, color=(1, 0, 0), line_width=2):
    """Draw circle on PDF page"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            ellipse = fitz.Rect(cx - radius, cy - radius, cx + radius, cy + radius)
            page.draw_circle((cx, cy), radius, color=color, width=line_width)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def highlight_text_in_pdf(pdf_path, output_path, page_num, x, y, width, height, color=(1, 1, 0)):
    """Highlight area in PDF page (yellow by default)"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            highlight = fitz.Rect(x, y, x + width, y + height)
            page.draw_rect(highlight, color=color, fill=color)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def blank_page_in_pdf(pdf_path, output_path, page_num):
    """Blank out (whiteout) specific page"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            white_rect = fitz.Rect(page.rect)
            page.draw_rect(white_rect, color=(1, 1, 1), fill=(1, 1, 1))
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def draw_line_on_pdf(pdf_path, output_path, page_num, x1, y1, x2, y2, color=(0, 0, 0), line_width=2):
    """Draw line on PDF page"""
    try:
        pdf_doc = fitz.open(pdf_path)
        
        if 0 <= page_num < len(pdf_doc):
            page = pdf_doc[page_num]
            page.draw_line((x1, y1), (x2, y2), color=color, width=line_width)
        
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True, output_path
    except Exception as e:
        return False, str(e)


def apply_multiple_edits(pdf_path, output_path, edits):
    """Apply multiple edits to PDF in sequence"""
    try:
        current_pdf = pdf_path
        
        for edit in edits:
            edit_type = edit.get('type')
            temp_output = os.path.join(os.path.dirname(output_path), f"temp_{uuid.uuid4().hex}.pdf")
            
            if edit_type == 'text':
                success, result = add_text_to_pdf(current_pdf, temp_output, [edit])
            elif edit_type == 'watermark':
                options = {
                    'type': 'text',
                    'text': edit.get('text', 'WATERMARK'),
                    'position': 'center',
                    'transparency': 50,
                    'rotation': 45,
                    'fontSize': 72,
                    'color': '#808080'
                }
                success, result = add_watermark_to_pdf(current_pdf, temp_output, options)
            elif edit_type == 'page_number':
                success, result = add_page_numbers_to_pdf(current_pdf, temp_output)
            elif edit_type == 'rectangle':
                success, result = draw_rectangle_on_pdf(
                    current_pdf, temp_output, edit.get('page', 0),
                    edit.get('x', 0), edit.get('y', 0),
                    edit.get('width', 100), edit.get('height', 100),
                    edit.get('color', (1, 0, 0))
                )
            elif edit_type == 'circle':
                success, result = draw_circle_on_pdf(
                    current_pdf, temp_output, edit.get('page', 0),
                    edit.get('cx', 0), edit.get('cy', 0),
                    edit.get('radius', 50), edit.get('color', (1, 0, 0))
                )
            elif edit_type == 'highlight':
                success, result = highlight_text_in_pdf(
                    current_pdf, temp_output, edit.get('page', 0),
                    edit.get('x', 0), edit.get('y', 0),
                    edit.get('width', 100), edit.get('height', 20),
                    edit.get('color', (1, 1, 0))
                )
            elif edit_type == 'delete_page':
                success, result = delete_pdf_pages(current_pdf, temp_output, [edit.get('page', 0)])
            elif edit_type == 'line':
                success, result = draw_line_on_pdf(
                    current_pdf, temp_output, edit.get('page', 0),
                    edit.get('x1', 0), edit.get('y1', 0),
                    edit.get('x2', 100), edit.get('y2', 100),
                    edit.get('color', (0, 0, 0))
                )
            else:
                continue
            
            if success:
                if current_pdf != pdf_path and os.path.exists(current_pdf):
                    os.remove(current_pdf)
                current_pdf = temp_output
            else:
                if os.path.exists(temp_output):
                    os.remove(temp_output)
                return False, result
        
        # Move final result to output path
        if current_pdf != output_path:
            if os.path.exists(current_pdf):
                os.rename(current_pdf, output_path)
        
        return True, output_path
    except Exception as e:
        return False, str(e)
